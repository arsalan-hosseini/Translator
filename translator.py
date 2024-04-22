import os
import sys

from docx import Document
from translate import Translator
from pdf2docx import parse

trans_list = []
lst_char = []
seq_lst = []


def pdf_to_docx(pdf, docx):
    parse(pdf, docx)


def translating(sequense):
    translator = Translator(to_lang='fa')
    for i in sequense:
        translation = translator.translate(i)
        trans_list.append(translation)
        print(translation)


def spacer(wordname, newname):
    list_of_char = []
    document = Document(wordname)
    new_document = Document()
    for para in document.paragraphs:
        for i in para.text:
            list_of_char.append(i)
            if i == '.':
                list_of_char.append('\n\n\n\n')

    print("The space was created!")
    sequense = ''.join(list_of_char)
    new_document.add_paragraph(sequense)
    new_document.save(newname)


def word_trans(wordname, new_word):
    document = Document(wordname)
    new_document = Document(new_word)
    for para in document.paragraphs:
        for i in para.text:
            lst_char.append(i)
            if i == '.':
               lst_char.append('\n')

    string = ''.join(lst_char)

    for i in string.split('\n'):
        if i:
            seq_lst.append(i)
    print("Sentences were extracted!")
    print("Translating...")
    translating(seq_lst)
    print("The sentences were translated!")
    print("Save information!")
    merging = [item for pair in zip(seq_lst, trans_list) for item in pair]
    mrg_n = [s + '\n' for  s in merging]
    mrg_str = ''.join(mrg_n)
    new_document.add_paragraph(mrg_str)
    new_document.save(new_word)


print("select option:\n\t1)pdf translator\n\t2)word translator\n\t3)pdf to word\n\t4)Create distance\n\t5)Exit")
while True:
    print("!@#$%^&*" * 20)
    select = input("enter number: ")
    if select == '1':
        print("Enter the names of the files with their extension.")
        pdf_name = input("pdf name: ")
        word_name = input("word name: ")
        word = 'sample.word'
        pdf_to_docx(pdf_name, word)
        print("pdf has been successfully converted to word!")
        word_trans(word, word_name)
        os.remove(word)
        print("Word file was created!")
    elif select == '2':
        print("Enter the names of the files with their extension.")
        word_name = input("Word name: ")
        new_word_name = input("New word name: ")
        word_trans(word_name, new_word_name)
        print("Word file was created!")
        print('OK!')
    elif select == '3':
        print("Enter the names of the files with their extension.")
        pdf_name = input("pdf name: ")
        word_name = input("word name: ")
        pdf_to_docx(pdf_name, word_name)
        print("the Word file was created!")
    elif select == '4':
        print("Enter the names of the files with their extension.")
        word_name = input("word name: ")
        new_word_name = input("new word name: ")
        spacer(word_name, new_word_name)
        print("Word file was created!")
    elif select == '5':
        sys.exit()
    else:
        print("Error! try again")
